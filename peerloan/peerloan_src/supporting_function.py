# -*- coding: utf-8 -*-
from django.contrib.sessions.models import Session
from django.core.exceptions import ObjectDoesNotExist
from django.db import transaction

from peerloan.models import Account
from peerloan.models import AuditTrail
from peerloan.models import AdminUser
from peerloan.models import BorrowRequest
from peerloan.models import Investment
from peerloan.models import Loan
from peerloan.models import LoanSchedule
from peerloan.models import Ledger
from peerloan.models import OTPMessage
from peerloan.models import Product
from peerloan.models import PendingTaskNotification
from peerloan.models import User
from peerloan.models import UserOperationRequest

import peerloan.peerloan_src.other_src.num2word as num2word
import peerloan.peerloan_class.admin_role as admin_role

from twilio.rest import TwilioRestClient

from datetime import datetime
import datetime as dt
from django.utils import timezone
import pytz
import json
import string
import random
import os
import csv
from docx import Document
import openpyxl as pyxl
from subprocess import Popen, PIPE
import re
import urllib

FLOAT_DATA_FORMAT = '{:,.2f}'

def get_lang(request):
	try:
		request.session['language']
	except KeyError:
		request.session['language'] = 'en'
	return request.session['language']

def get_user(request):
	sess = Session.objects.get(pk=request.COOKIES.get('sessionid'))
	if sess.get_decoded()['type'] == 'ADMIN':
		user = AdminUser.objects.get(email=sess.get_decoded()['email'])
	else:
		user = User.objects.get(email=sess.get_decoded()['email'], type=sess.get_decoded()['type'], status='ACTIVE')
	return user

def checkHTMLtags(ListOfInputs):
	pattern_1 = re.compile(r'.*<.*')
	pattern_2 = re.compile(r'.*>.*')
	for input in ListOfInputs:
		input=urllib.unquote(input)
		if pattern_1.match(input) or pattern_2.match(input):
			return True
	return False

def get_admin_role(usr_type):
	if usr_type == 'SUPERADMIN':
		return admin_role.SuperAdmin()
	elif usr_type == 'APPROVER':
		return admin_role.Approver()
	elif usr_type == 'OFFICER':
		return admin_role.Officer()
	elif usr_type == 'ACCOUNTANT':
		return admin_role.Accountant()
	else:
		return admin_role.AdminRole()

def date_postfix(n):
	if 10 < n and n <= 20:
		return 'th'
		
	if n % 10 == 1:
		return 'st'
	elif n % 10 == 2:
		return 'nd'
	elif n % 10 == 3:
		return 'rd'
	else:
		return 'th'
		
def try_KeyError(dict,key):
	try:
		return dict[key]
	except KeyError:
		return '--'

def import_public_holiday():
	# return a set of public holiday
	public_holidays = []
	f = open('/home/ubuntu/project_peerloan/peerloan/peerloan_src/public_holiday.csv', 'r')
	for row in csv.DictReader(f):
		public_holidays.append(row['Public Holiday'])
	f.close()
	return public_holidays

def  find_earliest_business_day(date,cut_off):
	"""
	inputs:
	date: datetime object
	cut_off: integer between 0 and 23
	
	output: datetime object
	"""
	public_holidays = import_public_holiday()
	while True:
		if date.strftime('%a') not in ['Sat', 'Sun'] and date.strftime('%d/%m/%Y') not in public_holidays:
			if date.hour < cut_off: # before 3pm
				break
			else:
				# set to 00:00:00 and find next business day
				date += dt.timedelta(days=1)
				date = datetime.strptime(date.strftime('%Y-%m-%d'), '%Y-%m-%d')
				date = pytz.timezone('Asia/Hong_Kong').localize(date)
		else: # find next business day
			date += dt.timedelta(days=1)
			date = datetime.strptime(date.strftime('%Y-%m-%d'), '%Y-%m-%d')
			date = pytz.timezone('Asia/Hong_Kong').localize(date)
	return date
	

def check_future_date_exists(date, months):
	"""
	date (datetime): starting date
	months (int): how many month has passed
	
	output: datetime object
	"""
	
	start_date_day = date.day
	start_date_month = (date.month + months) % 12
	if start_date_month == 0:
		start_date_month = 12
	start_date_year = date.year + ((date.month-1 + months) / 12)
	
	try:
		start_date = datetime.strptime(str(start_date_day)+'/'+str(start_date_month)+'/'+str(start_date_year), '%d/%m/%Y')
	except ValueError:
		start_date_day = 1
		start_date_month = (date.month + (months+1)) % 12
		if start_date_month == 0:
			start_date_month = 12
		start_date_year = date.year + ((date.month-1 + (months+1)) / 12)
		start_date = datetime.strptime('1/'+str(start_date_month)+'/'+str(start_date_year), '%d/%m/%Y') - dt.timedelta(days=1)
	return start_date

def generate_lender_no():
	this_year = str(timezone.localtime(timezone.now()).year)
	
	lender_no = str(len(User.objects.filter(type='L', create_timestamp__year=this_year))+1).zfill(6)
	lender_no = this_year[2] + lender_no[:2] + this_year[3] + lender_no[2:]
	return lender_no

def generate_ref_num(type, prefix):
	if type == 'borrow_request':
		today_date = timezone.localtime(timezone.now()).strftime('%Y/%m/%d')
		today_date = datetime.strptime(today_date, '%Y/%m/%d')
		last_bor = BorrowRequest.objects.filter(create_timestamp__startswith = today_date.date())
		if len(last_bor) >= 999:
			return HttpResponse('Numbers are all used.')
		if len(last_bor) == 0:
			ref_num = prefix + str(timezone.localtime(timezone.now()))[0:10].replace('-','') + '001'
		else:
			ref_num = prefix + str(timezone.localtime(timezone.now()))[0:10].replace('-','') + '%03d'%(len(last_bor)+1)
	if type in ['Deposit', 'Withdraw', 'Apply To Be Investor'] :
		today_date = timezone.localtime(timezone.now()).strftime('%Y/%m/%d')
		today_date = datetime.strptime(today_date, '%Y/%m/%d')
		last_trx = UserOperationRequest.objects.filter(create_timestamp__startswith = today_date.date())
		last_trx = [trx for trx in last_trx if type in trx.type]
		#last_trx = UserTransaction.objects.filter(type = type, create_timestamp__startswith = dt.date.today())
		if len(last_trx) >= 999:
			return HttpResponse('Numbers are all used.')
		if len(last_trx) == 0:
			ref_num = prefix + str(timezone.localtime(timezone.now()))[0:10].replace('-','') + '001'
		else:
			ref_num = prefix + str(timezone.localtime(timezone.now()))[0:10].replace('-','') + '%03d'%(len(last_trx)+1)
	return ref_num

def generate_unique_sequence(size=32, chars=string.ascii_uppercase + string.digits):
	path = '/home/ubuntu/project_peerloan/supporting_documents/'
	while 1:
		rand_seq = ''.join(random.choice(chars) for _ in range(size))
		existed_files = os.listdir(path)
		if rand_seq not in [existed_filename.split('.')[0] for existed_filename in existed_files]:
			return rand_seq

def calculate_loan_amount(inputs):
	major_scoring = {
	"Medical/Health": 26000,
	"Law": 15000,
	"Accounting": 12000,
	"Construction and Environment": 15000,
	"Engineering": 16000,
	"Design": 12000,
	"Business/Finance/Economic": 15000,
	"Education and Language": 12000,
	"Information Technology/Computing": 16000,
	"Social Sciences": 19000,
	"Hotel and Tourism": 12000,
	"Others": 10000,
	}
	resident_scoring = {
	"Public Housing": 0.6,
	"Owned by Family": 1,
	"Rent": 0.4,
	"Quarter": 0.8,
	"Student Hall of Residence": 0.4,
	}
	living_with_scoring = {
	"Parents": 1,
	"Relatives": 0.7,
	"Friends or Classmates": 0.4,
	"Others": 0.1,
	}
	university_scoring = {
	"The University of Hong Kong": 1,
	"The Chinese University of Hong Kong": 0.9,
	"The Hong Kong University of Science and Technology": 1,
	"The Hong Kong Polytechnic University": 0.8,
	"City University of Hong Kong": 0.9,
	"Hong Kong Baptist University": 0.9,
	"The Hong Kong Institute of Education": 1,
	"The Open University of Hong Kong": 0.4,
	"Lingnan University": 0.6,
	"Hong Kong Shue Yan University": 0.4,
	"Chu Hai College of Higher Education": 0.4,
	"Other Institutes": 0.3,
	}
	social_total_score_scoring = {
	'0-400': 0.95,
	'>400': 1
	}
	social_fraud_risk_score_scoring = {
	'0-1': 0.95,
	'>1': 1
	}
	
	major = inputs['major']
	resident = inputs['resident']
	living_with = inputs['living_with']
	university = inputs['university']
	GPA = float(inputs['GPA'])
	if GPA == None: # new student
		GPA = 2.0
	amount = major_scoring[major] * resident_scoring[resident] * living_with_scoring[living_with] * university_scoring[university] * float(GPA)
	
	# times social score factors
	social_total_score = inputs['social_total_score']
	social_fraud_risk_score = inputs['social_fraud_risk_score']
	amount *= social_total_score_scoring[social_total_score] * social_fraud_risk_score_scoring[social_fraud_risk_score]
	
	date_of_birth = datetime.strptime(inputs['date_of_birth'], '%Y-%m-%d')
	date_of_birth = pytz.timezone('Asia/Hong_Kong').localize(date_of_birth)
	age = (timezone.localtime(timezone.now()) - date_of_birth).days / 365
	
	# check the approval amount
	if 5000 <= amount and amount <= 7000:
		amount = 5000
	elif 7000 < amount and amount <= 12000:
		amount = 10000
	elif 12000 < amount and amount <= 17000:
		amount = 15000
	elif 17000 < amount and amount <= 22000:
		amount = 20000
	elif 22000 < amount and amount <= 27000:
		amount = 25000
	elif 27000 < amount and amount <= 32000:
		amount = 30000
	elif 32000 < amount and amount <= 37000:
		amount = 35000
	elif 37000 < amount:
		amount = 40000
	result = {'amount':amount, 'status':'ACCEPT'}
	
	# year 1, 2 only allow 20000
	if (inputs['year'] == 'Year 1' or inputs['year'] == 'Year 2') and amount > 20000:
		result['amount'] = 20000
		amount = result['amount']
	
	# check accept or not
	if amount < 5000:
		result['status'] = 'REJECT'
	if GPA < 1.5:
		result['status'] = 'REJECT'
	if age < 18 or age > 25:
		result['status'] = 'REJECT'
	
	if result['status'] == 'REJECT':
		result['amount'] = 0
	
	# check if is promotion product
	try:
		input['repayment_plan']
	except:
		''
	else:
		if input['repayment_plan'] == 'Promotion Balloon Payment':
			if amount < 7000:
				result['status'] = 'REJECT'
	
	return result

def bank_code_name_converter(code):
	f = open('/home/ubuntu/project_peerloan/peerloan/peerloan_src/bank_code_list.csv', 'r')
	for row in csv.DictReader(f):
	    if int(code) == int(row['Bank_Code']):
	    	return row['Bank_Name_in_English']
	f.close()
	return 'Unknown Bank Name'

@transaction.atomic
def check_OTP_match(inputs):
	"""
	inputs:
		usr id
		usr input OTP
		return link
		OTP action
	output:
		True or
		False + return link with error attribute
	"""
	usr_id = inputs['usr_id']
	usr_input_OTP = inputs['usr_input_OTP']
	return_link = inputs['return_link']
	OTP_action = inputs['OTP_action']
	result = True
	try:
		OTP_message_list = OTPMessage.objects.filter(usr_id=usr_id, action=OTP_action, status="ACTIVE")
		OTP_message = OTP_message_list.latest('generated_timestamp')
	except ObjectDoesNotExist:
		# receive OTP first
		"return redirect(return_link+'&error=invalid_OTP')"
		result, return_link = False, return_link+'&error=invalid_OTP'
		#result = True # remove when officially launch
	else:
		
		# check expiring
		if OTP_message.expiring_timestamp < timezone.localtime(timezone.now()):
			for item in OTP_message_list:
				item.status = 'EXPIRED'
				item.save()
			#return redirect(return_link+'&error=OTP_expired')
			result, return_link = False, return_link+'&error=OTP_expired'
		if usr_input_OTP != OTP_message.sequence:
			for item in OTP_message_list:
				item.status = 'EXPIRED'
				item.save()
			#return redirect(return_link+'&error=OTP_not_matched')
			result, return_link = False, return_link+'&error=OTP_not_matched'
		
	
	for item in OTP_message_list:
		item.status = 'EXPIRED'
		item.save()
	
	return result, return_link

def sync_to_usr(usr_id, dict):
	usr = User.objects.get(id=usr_id)
	details = json.loads(usr.detail_info)
	for k, v in dict.iteritems():
		details[k] = v
	usr.detail_info = json.dumps(details)
	usr.save()

def sync_from_usr(usr_id):
	usr = User.objects.get(id=usr_id)
	return json.loads(usr.detail_info)

def generate_portfolio_summary(usr_id, date):
	res_details = {
		'Portfolio Summary ($)': {},
		'Portfolio Summary (#)': {},
		'Return': {},
		'Net Flow and Delinquency ($)': {},
		'Net Flow and Delinquency (#)': {},
	}
	
	# update portfolio summary $ & #
	uor_list = UserOperationRequest.objects.filter(usr_id=usr_id, type='Deposit Money', status='CONFIRMED')
	# 1
	deposited_amount = 0 
	for uor in uor_list:
		if uor.create_timestamp.year == date.year and uor.create_timestamp.month == date.month:
			details = json.loads(uor.details)
			deposited_amount += float(details['transferred_amount'])
	# 2
	unallocated_amount = Account.objects.get(usr_id=usr_id).balance
	# 3
	inv_list = Investment.objects.filter(usr_id=usr_id)
	unmatched_amount = sum([inv.usable_amount+inv.on_hold_amount for inv in inv_list])
	# 4 5 6 7 8 9 10
	matched_amount = 0
	accrued_interest = 0
	normal_run_down_amount = 0
	normal_settlement_amount = 0
	early_settlement_amount = 0
	new_account_amt = 0
	gross_charge_off_amt = 0
	# 11 12 13 14 15
	active_account_no = 0
	new_account_no = 0
	normal_settlement_no = 0
	early_settlement_no = 0
	gross_charge_off_no = 0
	# 16 17 18
	weighted_annual_interest_rate = 0
	risk_adjusted_return = 0
	interest_received = 0
	# 19 20 21 22 23 24
	current_amt = 0
	cycle_1_amt = 0
	cycle_2_amt = 0
	cycle_3_amt = 0
	cycle_4_amt = 0
	annualized_loss_rate = 0
	# 25 26 27 28 29 30 31
	current_no = 0
	cycle_1_no = 0
	cycle_2_no = 0
	cycle_3_no = 0
	cycle_4_no = 0
	write_off_no = 0
	write_off_percentage = 0
	
	for inv in inv_list:
		loan_list = Loan.objects.filter(inv_id=inv.id)
		for loan in loan_list:
			try:
				bor = BorrowRequest.objects.get(id=loan.bor_id)
			except:
				continue
				
			if 'PAYBACK' not in bor.status or 'PAYBACK COMPLETED' in bor.status:
				continue
			if bor.draw_down_date != None:
				if bor.draw_down_date.year == date.year and bor.draw_down_date.month == date.month:
					# 9
					new_account_amt += loan.initial_amount
					# 12
					new_account_no += 1
				d_d_date = datetime.strptime(datetime.strftime(bor.draw_down_date,'%m/%Y'),'%m/%Y')
				e_e_date = datetime.strptime(datetime.strftime(bor.expected_end_date,'%m/%Y'),'%m/%Y')
				d_d_date = pytz.timezone('Asia/Hong_Kong').localize(d_d_date)
				e_e_date = pytz.timezone('Asia/Hong_Kong').localize(e_e_date)
				if date >= d_d_date and date <= e_e_date:
					# 11
					active_account_no += 1
					# 4
					matched_amount += loan.remain_principal_lender
			ratio = float(loan.initial_amount) / bor.amount
			
			los_list = LoanSchedule.objects.filter(bor_id=bor.id)
			filtered_los_list = [los for los in los_list if datetime.strptime(los.due_date,'%Y/%m/%d').year == date.year and 
			datetime.strptime(los.due_date,'%Y/%m/%d').month == date.month]
			# 5
			accrued_interest += sum([los.interest_l for los in filtered_los_list]) * ratio
			# 6 7 8 (some error in 7)
			normal_run_down_amount += sum([los.principal_l + los.interest_l for los in filtered_los_list if los.repayment_method=='Auto Pay' and 'PAID' in los.status]) * ratio
			normal_settlement_amount += sum([los.principal_l + los.interest_l for los in filtered_los_list if los.repayment_method=='Instalment' and 'PAID' in los.status]) * ratio
			early_settlement_amount += sum([los.principal_l + los.interest_l for los in filtered_los_list if los.repayment_method=='Early Instalment' and 'PAID' in los.status]) * ratio
			# 13 14
			normal_settlement_no += len([los for los in filtered_los_list if los.repayment_method=='Instalment' and 'PAID' in los.status])
			early_settlement_no += len([los for los in filtered_los_list if los.repayment_method=='Early Instalment' and 'PAID' in los.status])
			# 18
			interest_received += sum([los.interest_l for los in filtered_los_list if 'PAID' in los.status]) * ratio
			
			# 19 20 21 22 23
			current_amt += sum([los.principal_l + los.interest_l for los in filtered_los_list if los.overdue_days == 0 and 'PAID' in los.status]) * ratio
			cycle_1_amt += sum([los.principal_l + los.interest_l for los in filtered_los_list if los.overdue_days >= 1 and los.overdue_days <= 29 and 'PAID' in los.status]) * ratio
			cycle_2_amt += sum([los.principal_l + los.interest_l for los in filtered_los_list if los.overdue_days >= 30 and los.overdue_days <= 59 and 'PAID' in los.status]) * ratio
			cycle_3_amt += sum([los.principal_l + los.interest_l for los in filtered_los_list if los.overdue_days >= 60 and los.overdue_days <= 89 and 'PAID' in los.status]) * ratio
			cycle_4_amt += sum([los.principal_l + los.interest_l for los in filtered_los_list if los.overdue_days >= 90 and los.overdue_days <= 120 and 'PAID' in los.status]) * ratio
			
			# 25 26 27 28 29
			current_no += len([los for los in filtered_los_list if los.overdue_days == 0 and 'PAID' in los.status])
			cycle_1_no += len([los for los in filtered_los_list if los.overdue_days >= 1 and los.overdue_days <= 29 and 'PAID' in los.status])
			cycle_2_no += len([los for los in filtered_los_list if los.overdue_days >= 30 and los.overdue_days <= 59 and 'PAID' in los.status])
			cycle_3_no += len([los for los in filtered_los_list if los.overdue_days >= 60 and los.overdue_days <= 89 and 'PAID' in los.status])
			cycle_4_no += len([los for los in filtered_los_list if los.overdue_days >= 90 and los.overdue_days <= 120 and 'PAID' in los.status])
			
			if loan.update_timestamp.year == date.year and loan.update_timestamp.month == date.month and loan.status == 'WRITE OFF':
				# 10
				gross_charge_off_amt += loan.remain_principal_l
				# 15
				gross_charge_off_no += 1
				# 30
				write_off_no += 1
	# 16
	for inv in inv_list:
		loan_list = Loan.objects.filter(inv_id=inv.id)
		for loan in loan_list:
			try:
				bor = BorrowRequest.objects.get(id=loan.bor_id)
			except:
				continue
			prod = Product.objects.get(id=bor.prod_id)
			if bor.draw_down_date != None:
				d_d_date = datetime.strptime(datetime.strftime(bor.draw_down_date,'%m/%Y'),'%m/%Y')
				e_e_date = datetime.strptime(datetime.strftime(bor.expected_end_date,'%m/%Y'),'%m/%Y')
				d_d_date = pytz.timezone('Asia/Hong_Kong').localize(d_d_date)
				e_e_date = pytz.timezone('Asia/Hong_Kong').localize(e_e_date)
				if date >= d_d_date and date <= e_e_date:
					if matched_amount != 0:
						weighted_annual_interest_rate += (loan.remain_principal_lender/matched_amount) * prod.APR_lender
					else:
						weighted_annual_interest_rate = 0
	# 24
	if matched_amount != 0:
		annualized_loss_rate = gross_charge_off_amt*12/matched_amount
	# 31
	if (active_account_no + normal_settlement_no + early_settlement_no) != 0:
		write_off_percentage = write_off_no / (active_account_no + normal_settlement_no + early_settlement_no)
	# 17
	risk_adjusted_return = weighted_annual_interest_rate - write_off_percentage
	portfolio_summary_amt = {
		'Deposited Amount': round(deposited_amount, 2),
		'Unallocated Amount': round(unallocated_amount, 2),
		'Unmatched Amount': round(unmatched_amount, 2),
		'Matched Amount': round(matched_amount, 2),
		'Accrued Interest': round(accrued_interest, 2),
		'Normal Run-Down': round(normal_run_down_amount, 2),
		'Normal Settlement': round(normal_settlement_amount, 2),
		'Early Settlement': round(early_settlement_amount, 2),
		'New Account': round(new_account_amt, 2),
		'Gross Charge Off': round(gross_charge_off_amt, 2),
	}
	portfolio_summary_no = {
		'Active Account': active_account_no,
		'New Account': new_account_no,
		'Normal Settlement': normal_settlement_no,
		'Early Settlement': early_settlement_no,
		'Gross Charge Off': gross_charge_off_no,
	}
	return_ = {
		'Weighted Annual Interest Rate': weighted_annual_interest_rate,
		'Risk Adjusted Return': risk_adjusted_return,
		'Interest Received': interest_received
	}
	net_flow_and_delinquency_amt = {
		"0 DPD (Current)": current_amt,
		"1-29 DPD (Cycle 1)": cycle_1_amt,
		"30-59 DPD (Cycle 2)": cycle_2_amt,
		"60-89 DPD (Cycle 3)": cycle_3_amt,
		"90-120 DPD (Cycle 4)": cycle_4_amt,
		"Annualized Loss Rate": annualized_loss_rate
	}
	net_flow_and_delinquency_no = {
		"0 DPD (Current)": current_no,
		"1-29 DPD (Cycle 1)": cycle_1_no,
		"30-59 DPD (Cycle 2)": cycle_2_no,
		"60-89 DPD (Cycle 3)": cycle_3_no,
		"90-120 DPD (Cycle 4)": cycle_4_no,
		"Write-Off (#)": write_off_no,
		"Write-Off Accounts Rate": write_off_percentage
	}
	res_details['Portfolio Summary ($)'] = portfolio_summary_amt
	res_details['Portfolio Summary (#)'] = portfolio_summary_no
	res_details['Return'] = return_
	res_details['Net Flow and Delinquency ($)'] = net_flow_and_delinquency_amt
	res_details['Net Flow and Delinquency (#)'] = net_flow_and_delinquency_no
	
	return res_details

@transaction.atomic
def update_ledger(inputs):
	try:
		inputs['usr_id']
	except KeyError: # borrower ledger
		try:
			last_led = Ledger.objects.filter(bor_id=inputs['bor_id']).latest('id')
		except ObjectDoesNotExist: # fit to old data
			first_led = Ledger(
			bor_id = inputs['bor_id'],
			description = 'Begin Balance',
			reference = '--',
			debit = 0,
			credit = 0,
			balance = 0,
			create_timestamp = inputs['datetime'] - dt.timedelta(seconds=300),
			update_timestamp = inputs['datetime'] - dt.timedelta(seconds=300)
			)
			first_led.save()
			last_led = Ledger.objects.filter(bor_id=inputs['bor_id']).latest('id')
		
		new_led = Ledger(
		bor_id = inputs['bor_id'],
		description = inputs['description'],
		reference = inputs['reference'],
		debit = inputs['debit'],
		credit = inputs['credit'],
		balance = last_led.balance - inputs['debit'] + inputs['credit'],
		create_timestamp = inputs['datetime'],
		update_timestamp = inputs['datetime']
		)
		try:
			new_led.status = inputs['status']
		except KeyError:
			''
			
		new_led.save()
		pass

	else: # lender ledger
		# begin balance
		if inputs['description'] == 'Begin Balance':
			new_led = Ledger(
			usr_id = inputs['usr_id'],
			description = inputs['description'],
			reference = inputs['reference'],
			debit = inputs['debit'],
			credit = inputs['credit'],
			balance = 0,
			create_timestamp = inputs['datetime'],
			update_timestamp = inputs['datetime']
			)
			new_led.save()
			return None
			
		try:
			last_led = Ledger.objects.filter(usr_id=inputs['usr_id']).latest('id')
		except ObjectDoesNotExist: # fit to old data
			first_led = Ledger(
			usr_id = inputs['usr_id'],
			description = 'Begin Balance',
			reference = '--',
			debit = 0,
			credit = 0,
			balance = 0,
			create_timestamp = inputs['datetime'] - dt.timedelta(seconds=300),
			update_timestamp = inputs['datetime'] - dt.timedelta(seconds=300)
			)
			first_led.save()
			last_led = Ledger.objects.filter(usr_id=inputs['usr_id']).latest('id')
		
		new_led = Ledger(
		usr_id = inputs['usr_id'],
		description = inputs['description'],
		reference = inputs['reference'],
		debit = inputs['debit'],
		credit = inputs['credit'],
		balance = last_led.balance - inputs['debit'] + inputs['credit'],
		create_timestamp = inputs['datetime'],
		update_timestamp = inputs['datetime']
		)
		try:
			new_led.status = inputs['status']
		except KeyError:
			''
		new_led.save()
		pass

def update_aut(inputs):
	new_aut = AuditTrail(
	usr_id = inputs['usr_id'],
	description = inputs['description'],
	ref_id = inputs['ref_id'],
	model = inputs['model'],
	#details = json.dumps(inputs['details']),
	by = inputs['by'],
	create_timestamp = inputs['datetime'],
	update_timestamp = inputs['datetime']
	)
	try:
		new_aut.details = json.dumps(inputs['details'])
	except KeyError:
		''
	new_aut.save()
	pass

@transaction.atomic
def update_ptn(inputs):
	if inputs['action'] == 'create':
		new_ptn = PendingTaskNotification(
		ref_id = inputs['ref_id'],
		model = inputs['model'],
		type = inputs['type'],
		description = inputs['description'],
		status = inputs['status'],
		create_timestamp = inputs['datetime'],
		update_timestamp = inputs['datetime']
		)
		new_ptn.save()
		pass
	elif inputs['action'] == 'modify':
		try:
			ptn = PendingTaskNotification.objects.get(ref_id=inputs['ref_id'], model=inputs['model'])
		except ObjectDoesNotExist: # fit old data
			ptn = PendingTaskNotification(
			ref_id = inputs['ref_id'],
			model = inputs['model'],
			type = inputs['type'],
			description = inputs['description'],
			status = inputs['status'],
			create_timestamp = inputs['datetime'],
			update_timestamp = inputs['datetime']
			)
			ptn.save()
		else:
			ptn.description = inputs['description']
			ptn.status = inputs['status']
			ptn.type = inputs['type']
			ptn.update_timestamp = inputs['datetime']
			ptn.save()
		pass

def find_ref_num(id, model):
	res = {
		'ref_num': '--',
		'type': '--',
	}
	if model == 'UserOperationRequest':
		try:
			uor = UserOperationRequest.objects.get(id=id)
		except ObjectDoesNotExist:
			res['ref_num'] = 'Ref is deleted'
		else:
			if uor.type == 'Deposit Money':
				details = json.loads(uor.details)
				res['ref_num'] = details['ref_num']
				res['type'] = uor.type
			elif uor.type == 'Withdraw Money':
				details = json.loads(uor.details)
				res['ref_num'] = details['ref_num']
				res['type'] = uor.type
			elif uor.type == 'Apply To Be Investor':
				details = json.loads(uor.details)
				res['ref_num'] = details['Individual']['Application No']
				res['type'] = uor.type
	elif model == 'BorrowRequest':
		try:
			bor = BorrowRequest.objects.get(id=id)
		except ObjectDoesNotExist:
			res['ref_num'] = 'Ref is deleted'
		else:
			res['ref_num'] = bor.ref_num
			res['type'] = 'BorrowRequest'
	return res

def send_mobile_message(inputs):
	ACCOUNT_SID = "ACb7c366a200aa5190530e128f46561915"
	AUTH_TOKEN = "c9287069a0ddd57109273330a0f5e973"
	client = TwilioRestClient(ACCOUNT_SID, AUTH_TOKEN)
	
	message = client.messages.create(
		body=inputs['body'],  # Message body, if any
		to=inputs['to'],
		from_="+13177080486",
	)

def generate_docx(id, model, doc_type):
	if model == 'UserOperationRequest':
		uor = UserOperationRequest.objects.get(id=id)
		if doc_type == 'Lender Agreement':
			int2word = {
				1:'First', 2:'Second', 3:'Third', 4:'Fourth', 5:'Fifth',
				6:'Sixth', 7:'Seventh', 8:'Eighth', 9:'Ninth', 10:'Tenth',
				11:'Eleventh', 12:'Twelfth', 13:'Thirteenth', 14:'Fourteenth', 15:'Fifteenth',
				16:'Sixteenth', 17:'Seventeenth', 18:'Eighteenth', 19:'Nineteenth', 20:'Twentieth',
				21:'Twenty-First', 22:'Twenty-Second', 23:'Twenty-Third', 24:'Twenty-Fourth', 25:'Twenty-Fifth',
				26:'Twenty-Sixth', 27:'Twenty-Seventh', 28:'Twenty-Eighth', 29:'Twenty-Ninth', 30:'Thirtieth', 31:'Thirty-First',
			}
			details = json.loads(uor.details)
			try:
				date = datetime.strptime(details['Confirm Lender Agreement Date'], '%Y/%m/%d %H:%M:%S')
				date = pytz.timezone('Asia/Hong_Kong').localize(date)
			except KeyError:
				date = timezone.localtime(timezone.now())
			try:
				details['Corporate']
			except KeyError:
				type = 'Individual'
				#date = datetime.strptime(details['Confirm Lender Agreement Date'], '%Y/%m/%d')
				info_dict = {
					#'DAY': int2word[int(date.day)],
					'DAY': str(date.day)+date_postfix(date.day),
					'MONTH': date.strftime('%B'),
					'YEAR': str(date.year),
					'NAME': details['Individual']['Surname'] + ' ' + details['Individual']['Given Name'],
					'MY_ONE': details['Individual']['HKID'],
					'ADDRESS': details['Individual']['Residential Address'],
					'OTP': details['OTP'],
					
					'FULL_DATE': '%s-%s-%s'%(str(date.day), date.strftime('%B'), str(date.year)),
					'FULL_TIME': date.strftime('%H:%M:%S')
				}
			else:
				type = 'Corporate'
				#date = datetime.strptime(details['Confirm Lender Agreement Date'], '%Y/%m/%d')
				info_dict = {
					'DAY': str(date.day)+date_postfix(date.day),
					'MONTH': date.strftime('%B'),
					'YEAR': str(date.year),
					'NAME': details['Corporate']['Company Name'],
					'CR_NO': details['Corporate']['CR NO.'],
					'ADDRESS': details['Corporate']['Office Address'],
					'OTP': details['OTP'],
					'MY_ONE': details['Individual']['HKID'],
					
					'FULL_DATE': '%s-%s-%s'%(str(date.day), date.strftime('%B'), str(date.year)),
					'FULL_TIME': date.strftime('%H:%M:%S')
				}
			doc = Document('/home/ubuntu/project_peerloan/document_peerloan/Investor_Agreement_'+type+'.docx')
			
			# update header part
			cnt = 0
			for p in doc.paragraphs:
				for k, v in info_dict.iteritems():
					if k in p.text:
						inline = p.runs
						for i in range(len(inline)):
							if k in inline[i].text:
								new_text = inline[i].text.replace(k, v)
								inline[i].text = new_text
								
			path = '/home/ubuntu/project_peerloan/tmp_file_peerloan/'
			dirs = os.listdir(path)
			doc.save(path+str(len(dirs))+'.docx')
	
	elif model == 'BorrowRequest':
		bor = BorrowRequest.objects.get(id=id)
		if doc_type == 'Loan Agreement':
			doc = Document('/home/ubuntu/project_peerloan/document_peerloan/Loan_Agreement.docx')
			details = json.loads(bor.detail_info)
			prod = Product.objects.get(id=bor.prod_id)
			
			if prod.repayment_plan == 'Promotion Balloon Payment':
				start_month = 4
			else:
				start_month = 1
			end_month = prod.repayment_period
			start_date = check_future_date_exists(timezone.localtime(timezone.now()), months=start_month)
			end_date = check_future_date_exists(timezone.localtime(timezone.now()), months=end_month)
			
			if prod.repayment_plan == 'Instalment':
				last_instalment_amount = round(bor.instalment_borrower, 2)
			elif prod.repayment_plan == 'Balloon Payment' or prod.repayment_plan == 'Promotion Balloon Payment':
				last_instalment_amount = round(bor.instalment_borrower + bor.amount, 2)
			date = datetime.strptime(details['Confirm Loan Agreement Date'], '%Y/%m/%d %H:%M:%S')
			"""
			MY_ZERO -> Loan Agreement Date
			MY_ONE -> HKID
			MY_TWO -> Due Date with th
			MY_THREE -> Bank Name
			MY_FOUR -> Account Number
			MY_FIVE -> Loan Amount in Chinese
			MY_SIX -> APR / 12
			
			btw: on9 python-docx :P
			"""
			info_dict = {
				
				'MY_ZERO': datetime.strftime(bor.create_timestamp, '%d %B %Y'),
				'LOAN_DRAWDOWN_DATE': datetime.strftime(timezone.localtime(timezone.now()), '%d %B %Y'),
				'NAME': '%s %s' % (details['Surname'], details['Given Name']),
				'MY_ONE': details['HKID'],
				'ADDRESS': details['Residential Address'],
				'LOAN_AMOUNT': FLOAT_DATA_FORMAT.format(bor.amount),
				'APR': FLOAT_DATA_FORMAT.format(bor.getBorAPR(prod.APR_borrower if prod.fake_APR_borrower == None else prod.fake_APR_borrower))+'%',
				'MY_SIX': FLOAT_DATA_FORMAT.format(bor.getBorAPR(prod.APR_borrower if prod.fake_APR_borrower == None else prod.fake_APR_borrower)/12)+'%',
				'NUM_OF_INSTALMENTS': str(prod.repayment_period if prod.fake_repayment_period == None else prod.fake_repayment_period),
				'INSTALMENT_AMOUNT': FLOAT_DATA_FORMAT.format(bor.instalment_borrower),
				'MY_TWO': str(timezone.localtime(timezone.now()).day) + date_postfix(timezone.localtime(timezone.now()).day),
				'DUE_DATE': str(timezone.localtime(timezone.now()).day),
				'FIRST_INSTALMENT': start_date.strftime('%d %B %Y'),
				'FIRST_AMOUNT': FLOAT_DATA_FORMAT.format(bor.instalment_borrower),
				'LAST_INSTALMENT': end_date.strftime('%d %B %Y'),
				'LAST_AMOUNT': FLOAT_DATA_FORMAT.format(last_instalment_amount),
				'MY_THREE': details['Bank Code'] + ' ' + bank_code_name_converter(details['Bank Code']),
				'MY_FOUR': details['Account Number'],
				'MY_FIVE': 'Hong Kong Dollars ' + num2word.Number2Word().float_to_dollar_english(bor.amount),
				'OTP': details['OTP'],
				
				'FULL_DATE': '%s-%s-%s'%(str(date.day), date.strftime('%B'), str(date.year)),
				'FULL_TIME': date.strftime('%H:%M:%S')
			}
			
			
			# update table
			table = doc.tables[0]
			for row in table.rows:
				for cell in row.cells:
					for p in cell.paragraphs:
						#print p.text
						inline = p.runs
						for i in range(len(inline)):
							#print '--'+inline[i].text
							for k, v in info_dict.iteritems():
								if k in inline[i].text:
									new_text = inline[i].text.replace(k, v)
									inline[i].text = new_text
									
			# update paragraph
			for p in doc.paragraphs:
				for k, v in info_dict.iteritems():
					if k in p.text:
						inline = p.runs
						for i in range(len(inline)):
							if k in inline[i].text:
								new_text = inline[i].text.replace(k, v)
								inline[i].text = new_text
								
			path = '/home/ubuntu/project_peerloan/tmp_file_peerloan/'
			dirs = os.listdir(path)
			doc.save(path+str(len(dirs))+'.docx')
			
		elif doc_type == 'Memorandum Agreement':
			doc = Document('/home/ubuntu/project_peerloan/document_peerloan/Memorandum_of_Agreement.docx')
			details = json.loads(bor.detail_info)
			prod = Product.objects.get(id=bor.prod_id)
			
			memo_confirm_date = datetime.strptime(details['Confirm Memorandum Date'], '%Y/%m/%d %H:%M:%S')
			memo_confirm_date = pytz.timezone('Asia/Hong_Kong').localize(memo_confirm_date)
			
			if prod.repayment_plan == 'Promotion Balloon Payment':
				start_month = 4
			else:
				start_month = 1
			
			end_month = prod.repayment_period
			start_date = check_future_date_exists(timezone.localtime(timezone.now()), months=start_month)
			end_date = check_future_date_exists(memo_confirm_date, months=end_month)
			
			if prod.repayment_plan == 'Instalment':
				last_instalment_amount = round(bor.instalment_borrower, 2)
			elif prod.repayment_plan == 'Balloon Payment' or prod.repayment_plan == 'Promotion Balloon Payment':
				last_instalment_amount = round(bor.instalment_borrower + bor.amount, 2)
			agreement_confirm_date = datetime.strptime(details['Confirm Loan Agreement Date'], '%Y/%m/%d %H:%M:%S')
			"""
			MY_ZERO -> Loan Agreement Date
			MY_ONE -> HKID
			MY_TWO -> Due Date with th
			MY_THREE -> Bank Name
			MY_FOUR -> Account Number
			MY_FIVE -> Loan Amount in Chinese
			btw: on9 python-docx :P
			"""
			
			info_dict = {
				
				'MY_ZERO': agreement_confirm_date.strftime('%d %B %Y'),
				'LOAN_DRAWDOWN_DATE': datetime.strftime(memo_confirm_date, '%d %B %Y'),
				'NAME': '%s %s' % (details['Surname'], details['Given Name']),
				'MY_ONE': details['HKID'],
				'ADDRESS': details['Residential Address'],
				'LOAN_AMOUNT': FLOAT_DATA_FORMAT.format(bor.amount),
				'APR': str(round(bor.getBorAPR(prod.APR_borrower if prod.fake_APR_borrower == None else prod.fake_APR_borrower), 2))+'%',
				'NUM_OF_INSTALMENTS': str(prod.repayment_period if prod.fake_repayment_period == None else prod.fake_repayment_period),
				'INSTALMENT_AMOUNT': FLOAT_DATA_FORMAT.format(bor.instalment_borrower),
				'MY_TWO': str(memo_confirm_date.day) + date_postfix(memo_confirm_date.day),
				'DUE_DATE': str(memo_confirm_date.day),
				'FIRST_INSTALMENT': start_date.strftime('%d %B %Y'),
				'FIRST_AMOUNT': FLOAT_DATA_FORMAT.format(bor.instalment_borrower),
				'LAST_INSTALMENT': end_date.strftime('%d %B %Y'),
				'LAST_AMOUNT': FLOAT_DATA_FORMAT.format(last_instalment_amount),
				'MY_THREE': details['Bank Code'] + ' ' + bank_code_name_converter(details['Bank Code']),
				'MY_FOUR': details['Account Number'],
				'MY_FIVE': 'Hong Kong Dollars ' + num2word.Number2Word().float_to_dollar_english(bor.amount),
				'OTP': details['Memorandum Captcha'],
				
				'FULL_DATE': '%s-%s-%s'%(str(timezone.localtime(timezone.now()).day), timezone.localtime(timezone.now()).strftime('%B'), str(timezone.localtime(timezone.now()).year)),
				'FULL_TIME': timezone.localtime(timezone.now()).strftime('%H:%M:%S')
			}
			
			# update table
			table = doc.tables[0]
			for row in table.rows:
				for cell in row.cells:
					for p in cell.paragraphs:
						#print p.text
						inline = p.runs
						for i in range(len(inline)):
							#print '--'+inline[i].text
							for k, v in info_dict.iteritems():
								if k in inline[i].text:
									new_text = inline[i].text.replace(k, v)
									inline[i].text = new_text
									
			# update paragraph
			for p in doc.paragraphs:
				for k, v in info_dict.iteritems():
					if k in p.text:
						inline = p.runs
						for i in range(len(inline)):
							if k in inline[i].text:
								new_text = inline[i].text.replace(k, v)
								inline[i].text = new_text
								
			path = '/home/ubuntu/project_peerloan/tmp_file_peerloan/'
			dirs = os.listdir(path)
			doc.save(path+str(len(dirs))+'.docx')
			
		elif doc_type == 'DDA':
			doc = Document('/home/ubuntu/project_peerloan/document_peerloan/DDA_form.docx')
			details = json.loads(bor.detail_info)
			#prod = Product.objects.get(id=bor.prod_id)
			
			memo_confirm_date = datetime.strptime(details['Confirm Memorandum Date'], '%Y/%m/%d %H:%M:%S')
			memo_confirm_date = pytz.timezone('Asia/Hong_Kong').localize(memo_confirm_date)
			
			info_dict = {
				'DATETIME': datetime.strftime(memo_confirm_date, '%Y/%m/%d'),
				'BN': bank_code_name_converter(details['Bank Code']),
				'BOR_NAME': details['Surname']+' '+details['Given Name'],
				'MOBILE': details['Mobile'],
			}
			int2alpha = {
				1: 'A',
				2: 'B',
				3: 'C',
				4: 'D',
				5: 'E',
				6: 'F',
				7: 'G',
				8: 'H',
				9: 'I',
				10: 'J',
				11: 'K',
				12: 'L',
				13: 'M',
				14: 'N',
				15: 'O',
			}
			for i in range(len(details['Bank Code'])):
				info_dict['BC'+int2alpha[i+1]] = details['Bank Code'][i]
			for i in range(len(details['Account Number'])):
				info_dict['XX'+int2alpha[i+1]] = details['Account Number'][i]
			for i in range(len(bor.ref_num)):
				info_dict['LN'+int2alpha[i+1]] = bor.ref_num[i]
			
			
			# update table
			table = doc.tables[1]
			for row in table.rows:
				for cell in row.cells:
					for p in cell.paragraphs:
						inline = p.runs
						for i in range(len(inline)):
							#print '--'+inline[i].text
							for k, v in info_dict.iteritems():
								if k in inline[i].text:
									new_text = inline[i].text.replace(k, v)
									inline[i].text = new_text
									
			# update account number
			entry_list = [20, 21, 22, 23, 24, 25, 26, 28, 30, 33, 35, 38]
			for i in range(len(entry_list)):
				entry = entry_list[i]
				try:
					table.rows[1].cells[entry].paragraphs[0].runs[0].text = info_dict['XX'+int2alpha[i+1]]
				except KeyError:
					table.rows[1].cells[entry].paragraphs[0].runs[0].text = ''
			
			# update paragraph
			cnt = 0
			for p in doc.paragraphs:
				for k, v in info_dict.iteritems():
					if k in p.text:
						inline = p.runs
						for i in range(len(inline)):
							if k in inline[i].text:
								new_text = inline[i].text.replace(k, v)
								inline[i].text = new_text
								
				cnt += 1
				if cnt == 2:
					break
			
			path = '/home/ubuntu/project_peerloan/tmp_file_peerloan/'
			dirs = os.listdir(path)
			doc.save(path+str(len(dirs))+'.docx')
	
	# convert doc to pdf
	command = 'echo ubuntu|sudo -S doc2pdf ' + path+str(len(dirs))+'.docx'
	process = Popen(command, stdout=PIPE, stderr=PIPE, shell=True)
	out, err = process.communicate()
	errcode = process.returncode
	if errcode != 0:
		raise Exception(err)
	# add password to pdf
	if doc_type in ['Memorandum Agreement', 'Loan Agreement', 'DDA']:
		mobile = details['Mobile']
	elif doc_type in ['Lender Agreement']:
		mobile = details['Individual']['Mobile']
	command = 'echo ubuntu|sudo -S pdftk ' + path+str(len(dirs))+'.pdf' + ' output ' + path+str(len(dirs))+'_.pdf userpw ' + '"%s"'%(mobile)
	process = Popen(command, stdout=PIPE, stderr=PIPE, shell=True)
	out, err = process.communicate()
	errcode = process.returncode
	if errcode != 0:
		raise Exception(err)
	
	return path+str(len(dirs))+'_.pdf'

def generate_xlsx(id, model, doc_type):
	if model == 'BorrowRequest':
		bor = BorrowRequest.objects.get(id=id)
		if doc_type == 'Repayment Schedule':
			wb = pyxl.load_workbook('/home/ubuntu/project_peerloan/document_peerloan/Repayment_Schedule.xlsx')
			ws = wb.active

			w2c_dict = {
				'name_of_borrower': 'D3',
				'interest_rate': 'H3',
				'loan_no': 'D5',
				'loan_date': 'H5',
				'loan_amount': 'D7',
				'total_tenors': 'H7',
				'start_due_date': 'B11',
				'start_principal': 'D11',
				'start_interest': 'E11',
				'start_repayment_amount': 'F11',
				'start_balance_of_principal': 'H11',
				'total_principal': 'D24',
				'total_interest': 'E24',
				'total_repayment_amount': 'F24',
			}
			
			details = json.loads(bor.detail_info)
			prod = Product.objects.get(id=bor.prod_id)

			try:
				loan_date = details['Confirm Memorandum Date']
			except KeyError:
				try:
					loan_date = details['Confirm Loan Agreement Date']
				except KeyError:
					loan_date = None
					
			# get repayment table
			inputs = {
				'date_type': 'month only',
				'start_balance': bor.amount,
				'rate_per_month': bor.getBorAPR(prod.APR_borrower) * 0.01 / 12,
				'instalment': bor.instalment_borrower,
				'repayment_plan': prod.repayment_plan,
				'repayment_period': prod.repayment_period,
			}
				
			if loan_date != None:
				loan_date = datetime.strptime(loan_date, '%Y/%m/%d %H:%M:%S')
				inputs['date_type'] = 'exact datetime'
				#first_date = check_future_date_exists(loan_date, months=1)
				inputs['start_date'] = {
						'day': loan_date.day,
						'month': loan_date.month,
						'year': loan_date.year
					}
				str_loan_date = loan_date.strftime('%Y/%m/%d')
			else:
				inputs['date_type'] = 'month only'
				str_loan_date = '--'
			repayment_schedule = generate_repayment_schedule(inputs)
			#print repayment_schedule
			if prod.repayment_plan == 'Instalment':
				total_repayment_amount = round(bor.instalment_borrower, 2) * prod.repayment_period
			elif prod.repayment_plan == 'Balloon Payment':
				total_repayment_amount = round(bor.instalment_borrower * prod.repayment_period, 2) + bor.amount
			elif prod.repayment_plan == 'Promotion Balloon Payment':
				total_repayment_amount = round(bor.instalment_borrower * prod.fake_repayment_period, 2) + bor.amount
			
			ws[w2c_dict['name_of_borrower']] = details['Surname']+' '+details['Given Name']
			ws[w2c_dict['interest_rate']] = FLOAT_DATA_FORMAT.format(bor.getBorAPR(prod.APR_borrower if prod.fake_APR_borrower == None else prod.fake_APR_borrower))+'%'
			ws[w2c_dict['loan_no']] = bor.ref_num
			ws[w2c_dict['loan_date']] = str_loan_date
			ws[w2c_dict['loan_amount']] = FLOAT_DATA_FORMAT.format(bor.amount)
			
			if prod.repayment_plan == 'Promotion Balloon Payment':
				ws[w2c_dict['total_tenors']] = str(9)
			else:
				ws[w2c_dict['total_tenors']] = str(prod.repayment_period)
			
			ws[w2c_dict['total_principal']] = FLOAT_DATA_FORMAT.format(bor.amount)
			ws[w2c_dict['total_interest']] = FLOAT_DATA_FORMAT.format(total_repayment_amount-bor.amount)
			ws[w2c_dict['total_repayment_amount']] = FLOAT_DATA_FORMAT.format(total_repayment_amount)
			for i in range(len(repayment_schedule)):
				row = repayment_schedule[i]
				alphabet, idx = w2c_dict['start_due_date'][0], int(w2c_dict['start_due_date'][1:])
				ws[alphabet+str(idx+i)] = row['Date']
				alphabet, idx = w2c_dict['start_principal'][0], int(w2c_dict['start_principal'][1:])
				ws[alphabet+str(idx+i)] = row['Returning Principal']
				alphabet, idx = w2c_dict['start_interest'][0], int(w2c_dict['start_interest'][1:])
				ws[alphabet+str(idx+i)] = row['Returning Interest']
				alphabet, idx = w2c_dict['start_repayment_amount'][0], int(w2c_dict['start_repayment_amount'][1:])
				ws[alphabet+str(idx+i)] = row['Instalment']
				alphabet, idx = w2c_dict['start_balance_of_principal'][0], int(w2c_dict['start_balance_of_principal'][1:])
				ws[alphabet+str(idx+i)] = row['Remain Balance']
			path = '/home/ubuntu/project_peerloan/tmp_file_peerloan/'
			dirs = os.listdir(path)
			wb.save(path+str(len(dirs))+'.xlsx')
	command = 'echo ubuntu|sudo -S unoconv ' + path+str(len(dirs))+'.xlsx'
	process = Popen(command, stdout=PIPE, stderr=PIPE, shell=True)
	out, err = process.communicate()
	errcode = process.returncode
	if errcode != 0:
		raise Exception(err)
	
	# add password to pdf
	command = 'echo ubuntu|sudo -S pdftk ' + path+str(len(dirs))+'.pdf' + ' output ' + path+str(len(dirs))+'_.pdf userpw ' + '"%s"'%(details['Mobile'])
	process = Popen(command, stdout=PIPE, stderr=PIPE, shell=True)
	out, err = process.communicate()
	errcode = process.returncode
	if errcode != 0:
		raise Exception(err)
	
	return path+str(len(dirs))+'_.pdf'

def generate_repayment_schedule(inputs):
	"""
	input:
	date_type (string): "exact datetime" or "month only"
	start_date (dict):
		day(int), month(int), year(int)
	start_balance (float): remain balance
	rate_per_month (float): APR / 12
	instalment (float)
	repayment_plan (string): "Instalment", "Balloon Payment", ...
	repayment_period (int)
	
	output: a table of repayment schedule
	Tenor
	Date
	Instalment
	Returning Principal
	Returning Interest
	Remain Balance
	"""
	remain_balance = inputs['start_balance']
	rate_per_month = inputs['rate_per_month']
	instalment = inputs['instalment']
	
	if inputs['repayment_plan'] == 'Instalment':
		repayment_table = []
		for i in range(inputs['repayment_period']):
			interest = remain_balance * rate_per_month
			principal = instalment - interest
			remain_balance -= principal
			
			if inputs['date_type'] == 'exact datetime':
				start_date = datetime.strptime(str(inputs["start_date"]["day"])+'/'+str(inputs["start_date"]["month"])+'/'+str(inputs["start_date"]["year"]), '%d/%m/%Y')
				date = check_future_date_exists(start_date, months=(i+1))
				date = date.strftime('%Y/%m/%d')
				
			elif inputs['date_type'] == 'month only':
				date = 'Month ' + str(i + 1)
			
			row = {
			'Tenor': str(i+1),
			'Date': date,
			'Instalment': FLOAT_DATA_FORMAT.format(instalment),
			'Returning Principal': FLOAT_DATA_FORMAT.format(principal),
			'Returning Interest': FLOAT_DATA_FORMAT.format(interest),
			'Remain Balance': FLOAT_DATA_FORMAT.format(abs(remain_balance)),
			}
			repayment_table.append(row)
	elif inputs['repayment_plan'] == 'Balloon Payment':
		repayment_table = []
		interest = remain_balance * rate_per_month
		for i in range(inputs['repayment_period']):
			if inputs['date_type'] == 'exact datetime':
				start_date = datetime.strptime(str(inputs["start_date"]["day"])+'/'+str(inputs["start_date"]["month"])+'/'+str(inputs["start_date"]["year"]), '%d/%m/%Y')
				date = check_future_date_exists(start_date, months=(i+1))
				date = date.strftime('%Y/%m/%d')
				
			elif inputs['date_type'] == 'month only':
				date = 'Month ' + str(i + 1)
			
			row = {
				'Tenor': str(i+1),
				'Date': date,
				'Instalment': '%.2f'%(interest+remain_balance) if i+1 == inputs['repayment_period'] else '%.2f'%(interest),
				'Returning Interest': '%.2f'%(interest),
				'Returning Principal': '%.2f'%(remain_balance) if i+1 == inputs['repayment_period'] else '%.2f'%(0),
				'Remain Balance': '%.2f'%(abs(0)) if i+1 == inputs['repayment_period'] else '%.2f'%(remain_balance),
			}
			repayment_table.append(row)
	elif inputs['repayment_plan'] == 'Promotion Balloon Payment':
		repayment_table = []
		interest = remain_balance * rate_per_month
		for i in range(inputs['repayment_period']):
			if inputs['date_type'] == 'exact datetime':
				start_date = datetime.strptime(str(inputs["start_date"]["day"])+'/'+str(inputs["start_date"]["month"])+'/'+str(inputs["start_date"]["year"]), '%d/%m/%Y')
				date = check_future_date_exists(start_date, months=(i+1))
				date = date.strftime('%Y/%m/%d')
				
			elif inputs['date_type'] == 'month only':
				date = 'Month ' + str(i + 1)
			
			row = {
				'Tenor': str(i+1),
				'Date': date,
				'Instalment': '%.2f'%(interest+remain_balance) if i+1 == inputs['repayment_period'] else '%.2f'%(interest),
				'Returning Interest': '%.2f'%(interest),
				'Returning Principal': '%.2f'%(remain_balance) if i+1 == inputs['repayment_period'] else '%.2f'%(0),
				'Remain Balance': '%.2f'%(abs(0)) if i+1 == inputs['repayment_period'] else '%.2f'%(remain_balance),
			}
			if i+1 < 4:
				row['Instalment'] = '%.2f'%(0)
				row['Returning Interest'] = '%.2f'%(0)
			repayment_table.append(row)
	return repayment_table

def generate_ldr_instalment_repayment_schedule(inputs):
	"""
	input:
	date_type (string): "exact datetime" or "month only"
	start_date (dict):
		day(int), month(int), year(int)
	start_balance (float): remain balance
	rate_per_month_bor (float): APR_bor / 12
	rate_per_month_ldr (float): APR_ldr / 12
	instalment (float)
	repayment_plan (string): "Instalment", "Balloon Payment", ...
	repayment_period (int)
	
	output: a table of repayment schedule
	Tenor
	Date
	Instalment
	Returning Principal
	Returning Interest
	Remain Balance
	"""
	remain_balance = inputs['start_balance']
	rate_per_month_bor = inputs['rate_per_month_bor']
	rate_per_month_ldr = inputs['rate_per_month_ldr']
	instalment = inputs['instalment']
	
	if inputs['repayment_plan'] == 'Instalment':
		repayment_table = []
		for i in range(inputs['repayment_period']):
			interest = remain_balance * rate_per_month_bor
			principal = instalment - interest
			remain_balance -= principal
			
			interest_ldr = interest * float(rate_per_month_ldr * 12) / float(rate_per_month_bor * 12)
			instalment_ldr = principal + interest_ldr
			
			if inputs['date_type'] == 'exact datetime':
				start_date = datetime.strptime(str(inputs["start_date"]["day"])+'/'+str(inputs["start_date"]["month"])+'/'+str(inputs["start_date"]["year"]), '%d/%m/%Y')
				date = check_future_date_exists(start_date, months=(i+1))
				date = date.strftime('%Y/%m/%d')
				
			elif inputs['date_type'] == 'month only':
				date = 'Month ' + str(i + 1)
			
			row = {
			'Tenor': str(i+1),
			'Date': date,
			'Instalment': FLOAT_DATA_FORMAT.format(instalment_ldr),
			'Returning Principal': FLOAT_DATA_FORMAT.format(principal),
			'Returning Interest': FLOAT_DATA_FORMAT.format(interest_ldr),
			'Remain Balance': FLOAT_DATA_FORMAT.format(abs(remain_balance)),
			}
			repayment_table.append(row)
	elif inputs['repayment_plan'] == 'Balloon Payment' or inputs['repayment_plan'] == 'Promotion Balloon Payment':
		repayment_table = []
		interest = remain_balance * rate_per_month_ldr
		for i in range(inputs['repayment_period']):
			if inputs['date_type'] == 'exact datetime':
				start_date = datetime.strptime(str(inputs["start_date"]["day"])+'/'+str(inputs["start_date"]["month"])+'/'+str(inputs["start_date"]["year"]), '%d/%m/%Y')
				date = check_future_date_exists(start_date, months=(i+1))
				date = date.strftime('%Y/%m/%d')
				
			elif inputs['date_type'] == 'month only':
				date = 'Month ' + str(i + 1)
			
			row = {
				'Tenor': str(i+1),
				'Date': date,
				'Instalment': '%.2f'%(interest+remain_balance) if i+1 == inputs['repayment_period'] else '%.2f'%(interest),
				'Returning Interest': '%.2f'%(interest),
				'Returning Principal': '%.2f'%(remain_balance) if i+1 == inputs['repayment_period'] else '%.2f'%(0),
				'Remain Balance': '%.2f'%(abs(0)) if i+1 == inputs['repayment_period'] else '%.2f'%(remain_balance),
			}
			repayment_table.append(row)
	return repayment_table

def calculate_repay_amount(inputs):
	"""
	inputs: 
	bor_id
	type: instalment / early settlement
	date (optional): calculate the repayment amount at some date, default is today
	
	outputs:
	1 (instalment):
		instalment amount
		overdue interest
		overpay amount
		instalment_month
	2. (early settlement):
		early_settlement_amount
		principal
		interest
		overpay amount
	"""
	bor_id = inputs['bor_id']
	type = inputs['type']
	date = try_KeyError(inputs, 'date')
	
	if date == '--':
		date = timezone.localtime(timezone.now())
	else:
		date = pytz.timezone('Asia/Hong_Kong').localize(datetime.strptime(date, '%Y/%m/%d'))
		
	bor = BorrowRequest.objects.get(id=bor_id)
	if type == 'instalment':
		try:
			los = LoanSchedule.objects.get(bor_id=bor.id,tenor=bor.repaid_month+1)
		except ObjectDoesNotExist:
			outputs = {
				'instalment_amount': 0,
				'overdue_interest': 0,
				'overpay_amount': 0,
				'late_charge': 0,
				'instalment_month': '--',
			}
		else:
			if los.status == 'OPEN':
				instalment_amount = (los.principal - los.paid_principal + los.interest - los.paid_interest)
				
				overdue_days = (date - pytz.timezone('Asia/Hong_Kong').localize(datetime.strptime(los.due_date, '%Y/%m/%d'))).days
				overdue_days = max(0, overdue_days)
				loan_list = Loan.objects.filter(bor_id=bor.id)
				prod = Product.objects.get(id=bor.prod_id)
				rate_per_day = bor.getBorAPR(prod.APR_borrower) * 0.01 / 360
				#remain_principal = sum([loan.remain_principal for loan in loan_list])
				overdue_interest_remained = los.instalment * rate_per_day * (overdue_days - los.overdue_interest_paid_days)
				
				overdue_interest = overdue_interest_remained + los.overdue_interest_unpay_paid
				late_charge = max(0, los.late_charge - los.paid_late_charge)
				instalment_month = los.due_date
				
			elif los.status == 'OVERDUE':
				# possibly exist more than one overdue instalment
				los_list = LoanSchedule.objects.filter(bor_id=bor.id,status='OVERDUE')
				
				instalment_amount = 0
				overdue_interest = 0
				late_charge = 0
				instalment_month = ''
				loan_list = Loan.objects.filter(bor_id=bor.id)
				prod = Product.objects.get(id=bor.prod_id)
				for los in los_list:
					instalment_amount += (los.principal - los.paid_principal + los.interest - los.paid_interest)
					
					overdue_days = (date - pytz.timezone('Asia/Hong_Kong').localize(datetime.strptime(los.due_date, '%Y/%m/%d'))).days
					overdue_days = max(0, overdue_days)
					rate_per_day = bor.getBorAPR(prod.APR_borrower) * 0.01 / 360
					#remain_principal = sum([loan.remain_principal for loan in loan_list])
					overdue_interest_remained = los.instalment * rate_per_day * (overdue_days - los.overdue_interest_paid_days)
					
					overdue_interest += overdue_interest_remained + los.overdue_interest_unpay_paid
					late_charge += max(0, los.late_charge - los.paid_late_charge)
					instalment_month += los.due_date + ', '
				instalment_month = instalment_month[0:-2]
			outputs = {
				'instalment_amount': instalment_amount,
				'overdue_interest': overdue_interest,
				'late_charge': late_charge,
				'overpay_amount': bor.overpay_amount,
				'instalment_month': instalment_month,
			}
	elif type == 'early_settlement':
		loan_list = Loan.objects.filter(bor_id=bor.id)
		los_list = LoanSchedule.objects.filter(bor_id=bor.id)
		prod = Product.objects.get(id=bor.prod_id)
		if bor.repaid_month == 0:
			last_pay_date = datetime.strptime(bor.draw_down_date.strftime('%Y/%m/%d'), '%Y/%m/%d')
			last_pay_date = pytz.timezone('Asia/Hong_Kong').localize(last_pay_date)
		else:
			los = LoanSchedule.objects.get(bor_id=bor.id, tenor=bor.repaid_month)
			last_pay_date = datetime.strptime(los.due_date, '%Y/%m/%d')
			last_pay_date = pytz.timezone('Asia/Hong_Kong').localize(last_pay_date)
		days = (date - last_pay_date).days
		days = max(0, days)
		rate_per_day = bor.getBorAPR(prod.APR_borrower) * 0.01 / 360
		remain_principal = sum([loan.remain_principal for loan in loan_list])
		early_settlement_amount = remain_principal *  (1 + rate_per_day * days)
		
		outputs = {
			'early_settlement_amount': early_settlement_amount,
			'principal': remain_principal,
			'interest': remain_principal *  (rate_per_day * days),
			'overpay_amount': bor.overpay_amount,
			'late_charge': sum([max(0, los.late_charge - los.paid_late_charge) for los in los_list]),
		}
	return outputs