#!/usr/bin/env python
# encoding: utf-8
import os, django
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "project_peerloan.settings")
django.setup()

from peerloan.models import BorrowRequest
from peerloan.models import Product
#import peerloan.peerloan_src.other_src.num2chi as num2chi

import peerloan.models as model
import peerloan.emailing as emailing
import peerloan.smsing as smsing
import peerloan.classes as classes

from django.db import transaction
from django.db.models import Q
from django.conf import settings
from datetime import datetime
import datetime as dt
from django.utils import timezone
import pytz
import peerloan.peerloan_src.supporting_function as sup_fn

from docx import Document
from subprocess import Popen, PIPE
import os, sys
import captcha.helpers
import re

import requests
from pprint import pprint
import openpyxl as pyxl
import json
from twilio.rest import TwilioRestClient
from twilio import TwilioRestException
from django.conf import settings

import math
import hashlib

inputs = {
	'date_type': 'month only',
	'start_balance': 30000.0,
	'rate_per_month_bor': 13.8 * 0.01 / 12,
	'rate_per_month_ldr': 10.0 * 0.01 / 12,
	'instalment': 2690.79 ,
	'repayment_plan': 'Instalment',
	'repayment_period': 12,
}

def UpdateHashedPassword():
	usr_list = model.AdminUser.objects.all()
	for usr in usr_list:
		usr.password = hashlib.md5(usr.password).hexdigest()
		usr.save()

bor = model.BorrowRequest.objects.get(id=38)
bor_emailing = emailing.BorrowerEmail(bor=bor)
bor_emailing.copy_of_loan_agreement()

#bor = BorrowRequest.objects.get(id=27)
#bor_emailing = emailing.BorrowerEmail(bor=bor)
#bor_emailing.copy_of_loan_agreement()
#print calculate_repay_amount(inputs)

#sup_fn.generate_xlsx(id=20, model='BorrowRequest', doc_type='Repayment Schedule')

#UpdateHashedPassword()
#print hashlib.md5("123123ZZ").hexdigest()
#print sup_fn.generate_ldr_instalment_repayment_schedule(inputs)

"""
ACCOUNT_SID = "ACb7c366a200aa5190530e128f46561915"
AUTH_TOKEN = "c9287069a0ddd57109273330a0f5e973"
client = TwilioRestClient(ACCOUNT_SID, AUTH_TOKEN)
try:
	
	message = client.messages.create(
		body='Zwap: test message 3 of 3',  # Message body, if any
		to="+852" + str('98321819'),
		from_="+13177080486",
	)
except TwilioRestException as e:
	print e


url = 'http://192.168.0.140:8080/sign_up/'
data = {
	"email": "heard I don't need to be email format?",
	"register_password": "I am a password that has : ?{}~`\	;",
	"csrfmiddlewaretoken:": "cYK2pt4DSFiPPmzcjuCw7oDlzCnkKec0"
}
r = requests.post(url, data=data)
print(r.text)

FLOAT_DATA_FORMAT = '{:,.2f}'

fs = classes.FriendlyScore()
fs.getToken()
code, fs_usr_details = fs.get(endpoint='users/partner-id/%s/show'%(usr.id), params={})

def check_further_date_exists(date, months):
	
	start_date_day = date.day
	start_date_month = (date.month + months) % 12
	if start_date_month == 0:
		start_date_month = 12
	start_date_year = date.year + ((date.month-1 + months) / 12)
	
	try:
		start_date = datetime.strptime(str(start_date_day)+'/'+str(start_date_month)+'/'+str(start_date_year), '%d/%m/%Y')
	except ValueError:
		start_date_day = 1
		start_date_month = (date.month + (months+1)) % 12
		if start_date_month == 0:
			start_date_month = 12
		start_date_year = date.year + ((date.month-1 + (months+1)) / 12)
		start_date = datetime.strptime('1/'+str(start_date_month)+'/'+str(start_date_year), '%d/%m/%Y') - dt.timedelta(days=1)
	return start_date
"""
#start_date = datetime.strptime('%s/%s/%s'%(start_date_year,start_date_month,start_date_day),'%Y/%m/%d')
#start_date = start_date.strftime('%d %B %Y')
inputs = {
	'bor_id': 69,
	'type': 'early_settlement',
	'date': '2016/11/01',
}
#outputs = sup_fn.calculate_repay_amount(inputs)
#print outputs

#print check_further_date_exists(datetime.strptime('31/10/2017', '%d/%m/%Y'), months=1)
#sup_fn.generate_docx(id=62, model='BorrowRequest', doc_type='Loan Agreement')
#sup_fn.generate_docx(id=311, model='UserOperationRequest', doc_type='Lender Agreement')

#bor = model.BorrowRequest.objects.get(id=2)
#bor_emailing = emailing.BorrowerEmail(bor=bor)
#prod = model.Product.objects.get(id = bor.prod_id)
#print bor.getMonthlyFlatRate(prod.repayment_period, prod.repayment_plan)
#bor_emailing.autopay_authorization_required()
#bor_emailing.copy_of_loan_agreement()
"""
now = datetime.strftime(timezone.localtime(timezone.now()), '%Y-%m-%d %H:%M:%S')
#print str(timezone.localtime(timezone.now()))
recove_time = datetime.strptime(now, '%Y-%m-%d %H:%M:%S')
recove_time = pytz.timezone('Asia/Hong_Kong').localize(recove_time)

print (timezone.localtime(timezone.now())-recove_time).seconds
"""
def funct1():
	new_uri = model.UserInformation(
	id = 100,
	section = 'section100',
	details = 'details100',
	info_type = 'type100'
	)
	new_uri.save()

@transaction.atomic
def TransactionTest():
	
	funct1()
	
	
	
	new_uri2 = model.UserInformation(
	id = 101,
	section = 'section101',
	details = 'details101',
	info_type = 'type101'
	)
	new_uri2.save()
	a = {}
	b = a['test']
#TransactionTest()

class Number2Word:
	
	NUMBER_WORDS = {
		0 : "Zero",
		1 : "One",
		2 : "Two",
		3 : "Three",
		4 : "Four",
		5 : "Five",
		6 : "Six",
		7 : "Seven",
		8 : "Eight",
		9 : "Nine",
		10 : "Ten",
		11 : "Eleven",
		12 : "Twelve",
		13 : "Thirteen",
		14 : "Fourteen",
		15 : "Fifteen",
		16 : "Sixteen",
		17 : "Seventeen",
		18 : "Eighteen",
		19 : "Nineteen",
		20 : "Twenty",
		30 : "Thirty",
		40 : "Forty",
		50 : "Fifty",
		60 : "Sixty",
		70 : "Seventy",
		80 : "Eighty",
		90 : "Ninety"
	}
	
	def int_to_english(self, n):
		n = int(n)
		
		english_parts = []
		ones = int(round(n % 10))
		tens = int(round(n % 100))
		hundreds = math.floor(n / 100) % 10
		thousands = math.floor(n / 1000)
		
		if thousands:
			english_parts.append(self.int_to_english(thousands))
			english_parts.append('Thousand')
			if not hundreds and tens:
				english_parts.append('and')
		if hundreds:
			english_parts.append(self.NUMBER_WORDS[hundreds])
			english_parts.append('Hundred')
			if tens:
				english_parts.append('and')
		if tens:
			if tens < 20 or ones == 0:
				english_parts.append(self.NUMBER_WORDS[tens])
			else:
				english_parts.append(self.NUMBER_WORDS[tens - ones])
				english_parts.append(self.NUMBER_WORDS[ones])
		return ' '.join(english_parts)
	
	def cent_to_english(self, n):
		english_parts = []
		cents = round((n - int(n)) * 100)
		tenths = int(round(cents % 100))
		hundredths = int(round(cents % 10))
		
		if tenths:
			if tenths < 20 or hundredths == 0:
				english_parts.append(self.NUMBER_WORDS[tenths])
			else:
				english_parts.append(self.NUMBER_WORDS[tenths - hundredths])
				english_parts.append(self.NUMBER_WORDS[hundredths])
		else:
			english_parts.append(self.NUMBER_WORDS[tenths])
		return ' '.join(english_parts)
	
	def float_to_dollar_english(self, n):
		int_english = self.int_to_english(n)
		cent_english = self.cent_to_english(n)
		english_parts = [int_english, 'Dollars', 'and', cent_english, 'Cents']
		return ' '.join(english_parts)

#n2w = Number2Word()

#print n2w.float_to_dollar_english(39999.80)
"""
bor = model.BorrowRequest.objects.get(id=75)
bor_emailing = emailing.BorrowerEmail(bor=bor)
bor_emailing.copy_of_loan_agreement()
"""

"""
fs = classes.FriendlyScore()
fs.getToken()
code, user = fs.get(endpoint='users/id/13216/show')
print user
"""
"""
def int_to_english(n):
	
	NUMBER_WORDS = {
		0 : "Zero",
		1 : "One",
		2 : "Two",
		3 : "Three",
		4 : "Four",
		5 : "Five",
		6 : "Six",
		7 : "Seven",
		8 : "Eight",
		9 : "Nine",
		10 : "Ten",
		11 : "Eleven",
		12 : "Twelve",
		13 : "Thirteen",
		14 : "Fourteen",
		15 : "Fifteen",
		16 : "Sixteen",
		17 : "Seventeen",
		18 : "Eighteen",
		19 : "Nineteen",
		20 : "Twenty",
		30 : "Thirty",
		40 : "Forty",
		50 : "Fifty",
		60 : "Sixty",
		70 : "Seventy",
		80 : "Eighty",
		90 : "Ninety"
	}
	
	
	english_parts = []
	ones = int(round(n % 10))
	tens = int(round(n % 100))
	hundreds = math.floor(n / 100) % 10
	thousands = math.floor(n / 1000)
	
	if thousands:
		english_parts.append(NUMBER_WORDS[thousands])
		english_parts.append('Thousand')
		if not hundreds and tens:
			english_parts.append('and')
	if hundreds:
		english_parts.append(NUMBER_WORDS[hundreds])
		english_parts.append('Hundred')
		if tens:
			english_parts.append('and')
	if tens:
		if tens < 20 or ones == 0:
			english_parts.append(NUMBER_WORDS[tens])
		else:
			english_parts.append(NUMBER_WORDS[tens - ones])
			english_parts.append(NUMBER_WORDS[ones])
	
	english_parts.append('Dollars and')
	cents = (n - int(n)) * 100
	tenths = int(round(cents % 100))
	hundredths = int(round(cents % 10))
	
	if tenths:
		if tenths < 20 or hundredths == 0:
			english_parts.append(NUMBER_WORDS[tenths])
		else:
			english_parts.append(NUMBER_WORDS[tenths - hundredths])
			english_parts.append(NUMBER_WORDS[hundredths])
	english_parts.append('Cents')
	
	return ' '.join(english_parts)

print int_to_english(25000)
"""
"""
ldr_usr_list = model.User.objects.filter(type='L')
for usr in ldr_usr_list:
	usr.detail_info = usr.detail_info.replace('"HKID or Passport No."', '"HKID"')
	usr.save()
	
"""
"""
bor = model.BorrowRequest.objects.get(ref_num='LOA20161011001')
led = model.Ledger.objects.filter(bor_id=bor.id).latest('id')
print led.description
"""
"""
table = doc.tables[0]
for row in table.rows:
	for cell in row.cells:
		for p in cell.paragraphs:
			print p.text
			inline = p.runs
			for i in range(len(inline)):
				print '--'+inline[i].text
				for k, v in info_dict.iteritems():
					if k in inline[i].text:
						new_text = inline[i].text.replace(k, v)
						inline[i].text = new_text
		
for p in doc.paragraphs:
	for k, v in info_dict.iteritems():
		if k in p.text:
			inline = p.runs
			for i in range(len(inline)):
				if k in inline[i].text:
					new_text = inline[i].text.replace(k, v)
					inline[i].text = new_text
path = '/home/ubuntu/project_peerloan/tmp_file_peerloan/'
dirs = os.listdir(path)
doc.save(path+str(len(dirs))+'.docx')
# update header part

"""
"""
#this_month = datetime.strptime('08/2016', '%m/%Y')
#this_month = pytz.timezone('Asia/Hong_Kong').localize(this_month)
#details = sup_fn.generate_portfolio_summary(usr_id=3, date=this_month)
#print details
inputs = {
	'usr_id':3,
	'description':'Money Exchange',
	'reference' :'PLME20160831001',
	'debit': 0,
	'credit': 10000,
	'datetime': timezone.localtime(timezone.now())
}
sup_fn.update_ledger(inputs)
"""
"""

due_date = pytz.timezone('Asia/Hong_Kong').localize(due_date)

print '=',datetime.now()
loan = Loan.objects.get(id=1)
print '==',timezone.localtime(loan.update_timestamp)
#loan.update_timestamp = datetime.now()
loan.update_timestamp = timezone.localtime(timezone.now())
loan.save()

print '===',timezone.localtime(timezone.now()).day

"""